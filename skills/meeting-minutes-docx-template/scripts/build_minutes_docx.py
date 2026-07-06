import argparse
import json
import os
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def body_element(doc):
    return doc._body._element


def clear_body_keep_section(doc):
    body = body_element(doc)
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def append_clone(doc, proto):
    new_p = deepcopy(proto._p)
    body = body_element(doc)
    sect = body.sectPr
    if sect is not None:
        body.remove(sect)
        body.append(new_p)
        body.append(sect)
    else:
        body.append(new_p)
    return doc.paragraphs[-1]


def clear_runs(paragraph):
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)


def copy_run_format(src_run, dst_run, bold=None):
    if src_run is not None and src_run._r.rPr is not None:
        dst_run._r.get_or_add_rPr()
        dst_run._r.rPr.clear()
        for child in src_run._r.rPr:
            dst_run._r.rPr.append(deepcopy(child))
    if bold is not None:
        dst_run.bold = bold


def set_text_from_proto(paragraph, text, bold=None):
    src_run = paragraph.runs[0] if paragraph.runs else None
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    copy_run_format(src_run, run, bold=bold)


def set_time_text(paragraph, text):
    src_runs = list(paragraph.runs)
    clear_runs(paragraph)
    if text.startswith("时间"):
        label = "时间"
        rest_text = text[2:]
    else:
        label = "时间"
        rest_text = "：" + text
    prefix = paragraph.add_run(label)
    copy_run_format(src_runs[0] if src_runs else None, prefix, bold=True)
    rest = paragraph.add_run(rest_text)
    src = src_runs[1] if len(src_runs) > 1 else (src_runs[0] if src_runs else None)
    copy_run_format(src, rest, bold=False)


def set_answer_text(paragraph, text):
    src_runs = list(paragraph.runs)
    clear_runs(paragraph)
    label = paragraph.add_run("A：")
    copy_run_format(src_runs[0] if src_runs else None, label, bold=True)
    body_text = text[2:] if text.startswith("A：") else text
    body = paragraph.add_run(body_text)
    src = src_runs[1] if len(src_runs) > 1 else (src_runs[0] if src_runs else None)
    copy_run_format(src, body, bold=False)


def find_first(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    raise RuntimeError(f"prototype paragraph not found: {label}")


def paragraph_num_id(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
        return None
    return int(p_pr.numPr.numId.val)


def load_content(path):
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    required = ["title", "date", "summary", "company_overview", "qa"]
    missing = [k for k in required if k not in data]
    if missing:
        raise ValueError(f"content JSON missing keys: {missing}")
    return data


def collect_prototypes(template_doc):
    proto = {}
    proto["title"] = template_doc.paragraphs[0]
    proto["date"] = template_doc.paragraphs[1]
    proto["blank"] = next((p for p in template_doc.paragraphs if not p.text.strip()), template_doc.paragraphs[2])
    proto["summary_title"] = find_first(template_doc, lambda p: p.text.strip() == "要点总结：", "要点总结")
    proto["summary_topic"] = find_first(
        template_doc,
        lambda p: p.style.name == "List Paragraph" and paragraph_num_id(p) == 1 and p.text.strip(),
        "summary first-level topic",
    )
    detail_by_num_id = {}
    for p in template_doc.paragraphs:
        if p.style.name != "List Paragraph" or not p.text.strip():
            continue
        num_id = paragraph_num_id(p)
        if num_id is not None and num_id != 1 and num_id not in detail_by_num_id:
            detail_by_num_id[num_id] = p
    if not detail_by_num_id:
        raise RuntimeError("no summary detail numbering prototypes found")
    proto["summary_details"] = detail_by_num_id
    proto["summary_detail_fallback"] = next(iter(detail_by_num_id.values()))
    proto["company_title"] = find_first(template_doc, lambda p: p.text.strip() == "一、公司概要", "company overview heading")
    company_idx = next(i for i, p in enumerate(template_doc.paragraphs) if p.text.strip() == "一、公司概要")
    company_body = None
    for i, p in enumerate(template_doc.paragraphs):
        if i > company_idx and p.text.strip() and p.paragraph_format.first_line_indent is not None:
            company_body = p
            break
    if company_body is None:
        raise RuntimeError("prototype paragraph not found: company overview body")
    proto["company_body"] = company_body
    proto["qa_title"] = find_first(template_doc, lambda p: p.text.strip() == "二、Q&A", "Q&A heading")
    proto["qa_topic"] = find_first(template_doc, lambda p: re.match(r"^\d+\.\s+", p.text.strip()) is not None, "Q&A topic")
    proto["question"] = find_first(template_doc, lambda p: p.text.strip().startswith("Q："), "question")
    proto["answer"] = find_first(template_doc, lambda p: p.text.strip().startswith("A："), "answer")
    return proto


def build_docx(template_path, content_path, out_path):
    content = load_content(content_path)
    shutil.copyfile(template_path, out_path)
    template_doc = Document(template_path)
    doc = Document(out_path)
    proto = collect_prototypes(template_doc)
    clear_body_keep_section(doc)

    p = append_clone(doc, proto["title"])
    set_text_from_proto(p, content["title"], bold=True)
    p = append_clone(doc, proto["date"])
    set_time_text(p, content["date"])
    append_clone(doc, proto["blank"])
    p = append_clone(doc, proto["summary_title"])
    set_text_from_proto(p, "要点总结：", bold=True)

    for idx, block in enumerate(content["summary"], start=1):
        p = append_clone(doc, proto["summary_topic"])
        set_text_from_proto(p, block["topic"], bold=True)
        detail_proto = proto["summary_details"].get(idx + 1, proto["summary_detail_fallback"])
        for item in block.get("items", []):
            p = append_clone(doc, detail_proto)
            set_text_from_proto(p, item, bold=False)

    append_clone(doc, proto["blank"])
    p = append_clone(doc, proto["company_title"])
    set_text_from_proto(p, "一、公司概要", bold=True)
    for para in content["company_overview"]:
        p = append_clone(doc, proto["company_body"])
        set_text_from_proto(p, para, bold=False)

    append_clone(doc, proto["blank"])
    p = append_clone(doc, proto["qa_title"])
    set_text_from_proto(p, "二、Q&A", bold=True)
    total_qa_pairs = sum(len(block.get("items", [])) for block in content["qa"])
    qa_pair_index = 0
    for block in content["qa"]:
        p = append_clone(doc, proto["qa_topic"])
        set_text_from_proto(p, block["topic"], bold=True)
        for item in block.get("items", []):
            question = item["q"] if item["q"].startswith("Q：") else "Q：" + item["q"]
            answer = item["a"] if item["a"].startswith("A：") else "A：" + item["a"]
            p = append_clone(doc, proto["question"])
            set_text_from_proto(p, question, bold=True)
            p = append_clone(doc, proto["answer"])
            set_answer_text(p, answer)
            qa_pair_index += 1
            if qa_pair_index < total_qa_pairs:
                append_clone(doc, proto["blank"])

    doc.save(out_path)
    return validate_docx(out_path, content)


def all_runs_bold(paragraph):
    runs = [r for r in paragraph.runs if r.text]
    return bool(runs) and all(r.bold is True for r in runs)


def validate_docx(path, content):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    list_numbered = 0
    for p in doc.paragraphs:
        if p.style.name == "List Paragraph" and p._p.pPr is not None and p._p.pPr.numPr is not None:
            list_numbered += 1
    q_ok = True
    a_ok = True
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Q：") and not all_runs_bold(p):
            q_ok = False
        if t.startswith("A："):
            nonempty = [r for r in p.runs if r.text]
            if not nonempty or nonempty[0].bold is not True:
                a_ok = False
            if len(nonempty) > 1 and any(r.bold is True for r in nonempty[1:]):
                a_ok = False
    expected_summary_paras = sum(1 + len(b.get("items", [])) for b in content["summary"])
    expected_q = sum(len(b.get("items", [])) for b in content["qa"])
    qa_separator_ok = True
    seen_answers = 0
    in_qa = False
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        t = p.text.strip()
        if t == "二、Q&A":
            in_qa = True
            continue
        if in_qa and t.startswith("A："):
            seen_answers += 1
            if seen_answers < expected_q:
                if idx + 1 >= len(paragraphs) or paragraphs[idx + 1].text.strip() != "":
                    qa_separator_ok = False
    report = {
        "path": str(path),
        "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
        "list_numbered": list_numbered,
        "expected_summary_paragraphs": expected_summary_paras,
        "q_count": text.count("Q："),
        "expected_q_count": expected_q,
        "a_count": text.count("A："),
        "expected_a_count": expected_q,
        "tables": len(doc.tables),
        "q_bold_rule_ok": q_ok,
        "a_bold_rule_ok": a_ok,
        "qa_separator_rule_ok": qa_separator_ok,
    }
    return report


def main():
    parser = argparse.ArgumentParser(description="Build sell-side meeting minutes DOCX from a sample Word template and structured JSON content.")
    parser.add_argument("--template", required=True, help="Sample .docx template to copy and preserve")
    parser.add_argument("--content", required=True, help="UTF-8 JSON file with title/date/summary/company_overview/qa")
    parser.add_argument("--out", required=True, help="Output .docx path")
    parser.add_argument("--report", help="Optional JSON validation report path")
    args = parser.parse_args()

    template = Path(args.template)
    content = Path(args.content)
    out = Path(args.out)
    if not template.exists():
        raise FileNotFoundError(template)
    if not content.exists():
        raise FileNotFoundError(content)
    out.parent.mkdir(parents=True, exist_ok=True)

    report = build_docx(str(template), str(content), str(out))
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if args.report:
        with open(args.report, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()

